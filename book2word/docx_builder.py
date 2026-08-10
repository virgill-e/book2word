"""Génération du document Word : image nettoyée puis texte extrait, page par page."""
import io
import os
from dataclasses import dataclass, field
from typing import List, Optional, Sequence

import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


@dataclass
class PageContent:
    image: np.ndarray  # RGB
    texts: List[str] = field(default_factory=list)  # un élément par bloc de texte, dans l'ordre de lecture
    page_number: int = 0


def _encode_png(img: np.ndarray) -> bytes:
    ok, buf = cv2.imencode(".png", cv2.cvtColor(img, cv2.COLOR_RGB2BGR))
    if not ok:
        raise RuntimeError("Échec de l'encodage de l'image en PNG.")
    return buf.tobytes()


def _open_document(template_path: Optional[str]) -> Document:
    """Ouvre le document de base : un modèle (styles/police/mise en page) si fourni, sinon un
    document vierge. Un modèle réduit à un unique paragraphe vide (simple "porteur de style")
    voit ce paragraphe supprimé pour ne pas laisser de ligne blanche avant la première page."""
    if not template_path:
        return Document()
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Modèle introuvable : {template_path}")

    doc = Document(template_path)
    if len(doc.paragraphs) == 1 and not doc.tables and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    return doc


def build_docx(
    pages: Sequence[PageContent],
    output_path: str,
    max_image_width_in: float = 6.5,
    template_path: Optional[str] = None,
) -> None:
    """Construit le .docx : pour chaque page, l'image nettoyée puis le texte extrait en dessous.

    `template_path`, si fourni, sert de base (styles, police, mise en page) — le texte ajouté
    hérite du style "Normal" du modèle au lieu de la police par défaut de python-docx.
    """
    doc = _open_document(template_path)

    for i, page in enumerate(pages):
        png_bytes = _encode_png(page.image)
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(max_image_width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if page.texts:
            for block_text in page.texts:
                for line in block_text.split("\n"):
                    doc.add_paragraph(line)
                doc.add_paragraph("")
        else:
            p = doc.add_paragraph("[Aucun texte détecté sur cette page]")
            p.runs[0].italic = True

        if i != len(pages) - 1:
            doc.add_page_break()

    doc.save(output_path)
